"""Build the blinded candidate registries offered to reviewers, as one .docx.

    python paper/make_reviewer_supplement.py   # -> paper/build/reviewer_supplement.docx

Section 4.1.2 names ``data/h2_studies.json`` and ``data/h2_studies_v2.json`` as the
record of every candidate screened and the criterion each one failed. A reader of
the published paper reaches them through the repository. A reviewer cannot: Medical
Physics has been double-anonymised since 1 July 2026 and the repository URL names
the author, so the manuscript withholds it. This file closes that gap.

Two constraints shape the output. The submission system rejects ``.json`` for
supplementary material -- doc, docx, tex, pdf, ps, images, tds, xls and rtf are the
formats it accepts -- so the registries are rendered into a document rather than
attached as they are. And supporting material must itself be blinded, so this
refuses to write a file in which any identifying string survives. The check runs on
the rendered text, not on the input, because a substitution that silently matched
nothing is exactly the failure worth catching.

Every field of every record is emitted. Nothing is summarised or selected: the
rendering changes the container, not the content.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

PAPER = Path(__file__).resolve().parent
REPO = PAPER.parent
DEFAULT_OUTPUT = PAPER / "build" / "reviewer_supplement.docx"

SOURCES = (
    (
        REPO / "data" / "h2_studies.json",
        "First campaign registry",
        "The pool the paper reports. Section 4.1.2 and Appendix A.",
    ),
    (
        REPO / "data" / "h2_studies_v2.json",
        "Second campaign registry",
        "Nine candidates judged on full text, none admitted. Section 4.1.2.",
    ),
)

#: Every way the author's identity is spelled anywhere in this repository. The
#: replacement is deliberately not a blank: a reviewer should see that something
#: was withheld and why, rather than read a sentence with a hole in it.
IDENTIFYING = (
    (r"Yamamoto", "the author"),
    (r"Shuji", "the author"),
    (r"LISIT(?: Co\.,? Ltd\.?)?", "[affiliation withheld]"),
    (r"Institute[ -]of[ -]One", "[affiliation withheld]"),
    (r"TexelCraft(?: O[UÜ])?", "[affiliation withheld]"),
    (r"0000-0001-9211-1071", "[ORCID withheld]"),
    (r"[\w.+-]+@lisit\.jp", "[email withheld]"),
    (r"IORN-\d+[A-Z]?", "[project code withheld]"),
)

PREAMBLE = (
    "These are the two registries the manuscript refers to in Section 4.1.2 and "
    "in Appendix A. They record every study considered for the external "
    "validation of H2, in both campaigns, together with the criterion each "
    "excluded study failed and the page or table each admitted value was read "
    "from.",
    "They are supplied here because the manuscript withholds the repository URL "
    "and the archive DOI, both of which identify the author, and this is the "
    "material a reviewer would otherwise reach through them. The registries are "
    "machine-readable JSON in the repository; the submission system does not "
    "accept that format, so they are rendered here field by field. Every field "
    "of every record is present and none has been reordered or summarised.",
    "Strings that identify the author have been replaced; nothing else has been "
    "altered, and the full files are in the public repository disclosed at "
    "acceptance. The published papers in the pool are not redistributed, and "
    "these registries contain no imaging data of any kind.",
)


def blind(text: str) -> str:
    for pattern, replacement in IDENTIFYING:
        text = re.sub(pattern, replacement, text)
    return text


def surviving(text: str) -> list[str]:
    found: list[str] = []
    for pattern, _ in IDENTIFYING:
        found.extend(re.findall(pattern, text))
    return found


def _label(key: str) -> str:
    return key.replace("_", " ")


def _render(document, value, key, depth):
    """Emit one JSON node, whatever its shape, without dropping anything."""
    from docx.shared import Inches, Pt  # noqa: PLC0415

    if isinstance(value, dict):
        heading = document.add_paragraph()
        run = heading.add_run(_label(key))
        run.bold = True
        heading.paragraph_format.left_indent = Inches(0.25 * depth)
        heading.paragraph_format.space_before = Pt(6)
        for child_key, child in value.items():
            _render(document, child, child_key, depth + 1)
        return

    if isinstance(value, list):
        if value and isinstance(value[0], dict):
            heading = document.add_paragraph()
            run = heading.add_run(f"{_label(key)} ({len(value)})")
            run.bold = True
            heading.paragraph_format.left_indent = Inches(0.25 * depth)
            heading.paragraph_format.space_before = Pt(10)
            for index, item in enumerate(value, 1):
                marker = document.add_paragraph()
                marker.add_run(
                    f"{_label(key)} {index} of {len(value)}"
                ).italic = True
                marker.paragraph_format.left_indent = Inches(0.25 * (depth + 1))
                marker.paragraph_format.space_before = Pt(8)
                for child_key, child in item.items():
                    _render(document, child, child_key, depth + 2)
            return
        rendered = "; ".join(str(item) for item in value) if value else "(none)"
        value = rendered

    paragraph = document.add_paragraph()
    paragraph.add_run(f"{_label(key)}: ").bold = True
    paragraph.add_run("(none)" if value is None else str(value))
    paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
    paragraph.paragraph_format.space_after = Pt(2)


def build(output: Path = DEFAULT_OUTPUT) -> int:
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)

    title = document.add_paragraph()
    run = title.add_run(
        "Supplementary material for review: candidate study registries"
    )
    run.bold = True
    run.font.size = Pt(14)

    for paragraph in PREAMBLE:
        document.add_paragraph(paragraph)

    for source, heading, subtitle in SOURCES:
        if not source.is_file():
            raise SystemExit(f"{source} is missing")
        document.add_page_break()
        head = document.add_paragraph()
        head.add_run(heading).bold = True
        head.runs[0].font.size = Pt(13)
        note = document.add_paragraph()
        note.add_run(subtitle).italic = True

        payload = _blind_in_place(json.loads(source.read_text(encoding="utf-8")))
        for key, value in payload.items():
            _render(document, value, key, 0)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)

    rendered = "\n".join(p.text for p in Document(output).paragraphs)
    blinded_pattern_hits = surviving(rendered)
    if blinded_pattern_hits:
        output.unlink()
        raise SystemExit(
            f"the rendered document still names {sorted(set(blinded_pattern_hits))}; "
            "it has been deleted rather than left for upload"
        )

    print(f"wrote {output} ({output.stat().st_size // 1024} KB)")
    print(f"  {len(Document(output).paragraphs)} paragraphs, no identifying string")
    return 0


def _blind_in_place(payload):
    """Blind every string in the loaded JSON before it is rendered."""
    if isinstance(payload, dict):
        return {key: _blind_in_place(value) for key, value in payload.items()}
    if isinstance(payload, list):
        return [_blind_in_place(item) for item in payload]
    if isinstance(payload, str):
        return blind(payload)
    return payload


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return build(parser.parse_args(argv).output)


if __name__ == "__main__":
    raise SystemExit(main())
