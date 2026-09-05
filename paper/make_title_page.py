"""Build the separate title page Medical Physics requires as its own upload.

    python paper/make_title_page.py        # -> paper/build/title_page.docx

Medical Physics will not let the Files step proceed without it, and it must carry
the author identity that the manuscript itself may or may not. Everything here is
read from the manuscript's own front matter and from results/release.json, so the
title on the title page cannot drift from the title in the paper -- which is the
failure this file exists to prevent, the same class as a cover letter that keeps a
withdrawn claim.

The affiliation is the string Crossref carries for the author's published work,
city before postcode. Two orderings of one address read to an affiliation-matching
system as two affiliations.
"""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

PAPER = Path(__file__).resolve().parent
REPO = PAPER.parent
RENDERED = PAPER / "manuscript.md"
RELEASE = REPO / "results" / "release.json"
DEFAULT_OUTPUT = PAPER / "build" / "title_page.docx"

AUTHOR = "Shuji Yamamoto"
AFFILIATION = "Institute of One, LISIT Co., Ltd., Tokyo 150-0044, Japan"
EMAIL = "yamamoto@lisit.jp"
ORCID = "0000-0001-9211-1071"
ARTICLE_TYPE = "Research Article"


def read_title() -> str:
    text = RENDERED.read_text(encoding="utf-8")
    match = re.search(r'^title:\s*"(.+)"\s*$', text, re.M)
    if match is None:
        raise SystemExit("the rendered manuscript has no title in its front matter")
    return match.group(1)


def read_keywords() -> str:
    text = RENDERED.read_text(encoding="utf-8")
    match = re.search(r"^\*\*Keywords[:.]?\*\*\s*(.+?)\n\n", text, re.S | re.M)
    return " ".join(match.group(1).split()) if match else ""


def build(output: Path = DEFAULT_OUTPUT) -> int:
    from docx import Document  # noqa: PLC0415
    from docx.shared import Pt  # noqa: PLC0415

    release = json.loads(RELEASE.read_text(encoding="utf-8"))
    title = read_title()

    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(10)

    heading = document.add_paragraph()
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(14)

    document.add_paragraph(AUTHOR)
    document.add_paragraph(AFFILIATION)
    document.add_paragraph(f"ORCID: {ORCID}")

    document.add_paragraph()
    corresponding = document.add_paragraph()
    corresponding.add_run("Corresponding author. ").bold = True
    corresponding.add_run(f"{AUTHOR}, {AFFILIATION}. Email: {EMAIL}")

    document.add_paragraph()
    kind = document.add_paragraph()
    kind.add_run("Article type. ").bold = True
    kind.add_run(ARTICLE_TYPE)

    keywords = read_keywords()
    if keywords:
        line = document.add_paragraph()
        line.add_run("Keywords. ").bold = True
        line.add_run(keywords)

    document.add_paragraph()
    conflict = document.add_paragraph()
    conflict.add_run("Conflict of interest. ").bold = True
    conflict.add_run(
        "The author is Representative Director of LISIT Co., Ltd., which funds "
        "Institute of One and is his affiliation on this paper, and Chief Executive "
        "Officer of TexelCraft OU (Estonia). No product of either company is "
        "evaluated, recommended or used in this work. No other competing interest, "
        "financial or personal, is declared."
    )

    funding = document.add_paragraph()
    funding.add_run("Funding. ").bold = True
    funding.add_run("This research received no external funding.")

    data = document.add_paragraph()
    data.add_run("Data availability. ").bold = True
    data.add_run(
        f"All code, every results file the manuscript quotes and the pre-registration "
        f"documents are at {release['repository']}, released as "
        f"{release['version_tag']} (commit {release['release_commit'][:7]}) and "
        f"archived at https://doi.org/{release['zenodo_version_doi']}."
    )

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)
    print(f"wrote {output} ({output.stat().st_size // 1024} KB)")
    print(f"  title on the page matches the manuscript: {title[:60]}...")
    return 0


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    return build(parser.parse_args(argv).output)


if __name__ == "__main__":
    raise SystemExit(main())
